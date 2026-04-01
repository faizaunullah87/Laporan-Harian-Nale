import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import datetime

# --- FUNGSI BARU: Mengubah format tanggal menjadi teks Indonesia ---
def format_tanggal_indo(tgl):
    bulan_indo = ["", "Januari", "Februari", "Maret", "April", "Mei", "Juni", 
                  "Juli", "Agustus", "September", "Oktober", "November", "Desember"]
    return f"{tgl.day} {bulan_indo[tgl.month]} {tgl.year}"

def generate_report(template_file, tgl_indo, lokasi, kegiatan_data, kendala_data, followup_data, ip_photos):
    try:
        doc = Document(template_file)
    except Exception as e:
        st.error(f"Error: File '{template_file}' tidak ditemukan di folder!")
        return None

    # 1. INFORMASI UMUM (Tabel Index 0)
    # Menggunakan .add_run() agar teks ditambahkan di SAMPING tulisan yang sudah ada
    # Catatan: Jika di template tulisan "Tanggal" ada di kolom lain, kamu bisa ubah index [0, 3] ke kolom yang sesuai.
    cell_tgl = doc.tables[0].cell(0, 3).paragraphs[0]
    cell_tgl.add_run(" " + tgl_indo if cell_tgl.text else tgl_indo)

    cell_lokasi = doc.tables[0].cell(1, 3).paragraphs[0]
    cell_lokasi.add_run(" " + lokasi if cell_lokasi.text else lokasi)

    # 2. KEGIATAN HARIAN (Tabel Index 1)
    tabel_kegiatan = doc.tables[1]
    for idx, k in enumerate(kegiatan_data):
        row_index = idx + 1
        if row_index < len(tabel_kegiatan.rows):
            row = tabel_kegiatan.rows[row_index].cells
        else:
            row = tabel_kegiatan.add_row().cells
        row[0].text = str(idx + 1)
        row[1].text = k['waktu']
        row[2].text = k['uraian']
        row[3].text = k['ket']
        row[4].text = k['status']

    # 3. KENDALA / INSIDEN (Tabel Index 2)
    tabel_kendala = doc.tables[2]
    for idx, k in enumerate(kendala_data):
        row_index = idx + 1
        if row_index < len(tabel_kendala.rows):
            row = tabel_kendala.rows[row_index].cells
        else:
            row = tabel_kendala.add_row().cells
        row[0].text = str(idx + 1)
        row[1].text = k['masalah']
        row[2].text = k['tindakan']
        row[3].text = k['hasil']

    # 4. FOLLOW UP (Tabel Index 3)
    tabel_followup = doc.tables[3]
    for idx, f in enumerate(followup_data):
        row_index = idx + 1
        if row_index < len(tabel_followup.rows):
            row = tabel_followup.rows[row_index].cells
        else:
            row = tabel_followup.add_row().cells
        row[0].text = str(idx + 1)
        row[1].text = f['deskripsi']
        row[2].text = f['alasan']
        row[3].text = f['target']

    # 5. LAMPIRAN FOTO (Tabel paling bawah)
    tabel_foto = doc.tables[-1] 
    
    for row in tabel_foto.rows:
        for cell in row.cells:
            cell.text = ""
            
    valid_photos = [p for p in ip_photos if p['file'] is not None]
    
    row_idx = 0
    col_idx = 0
    for p in valid_photos:
        if row_idx >= len(tabel_foto.rows):
            tabel_foto.add_row()
            
        cell = tabel_foto.cell(row_idx, col_idx)
        
        para_img = cell.paragraphs[0]
        para_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = para_img.add_run()
        run_img.add_picture(p['file'], width=Inches(2.9)) 
        
        para_text = cell.add_paragraph()
        para_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_text.add_run(f"Status IP {p['name']} {p['status']}")
        
        col_idx += 1
        if col_idx > 1:
            col_idx = 0
            row_idx += 1

    target = BytesIO()
    doc.save(target)
    return target

# === UI STREAMLIT ===
st.set_page_config(page_title="Nale Report App", layout="wide")
st.title("📊 Auto Report Generator")

with st.sidebar:
    st.header("⚙️ Pengaturan Template")
    pilihan_template = st.radio(
        "Pilih Warna Kop Surat:",
        ["Template Biru", "Template Orange"]
    )
    
    file_template = "template_biru.docx" if pilihan_template == "Template Biru" else "template_orange.docx"

    st.divider()
    st.header("📝 Informasi Umum")
    # Input tanggal dari user
    tgl_input = st.date_input("Tanggal Report", datetime.date.today())
    lok = st.text_input("Lokasi / Site", "UNEJ")
    
    # Mengubah format inputan kalender menjadi format Indonesia (ex: 1 April 2025)
    teks_tgl_indo = format_tanggal_indo(tgl_input)
    st.info(f"Tanggal di laporan: {teks_tgl_indo}")

# --- FORM KEGIATAN HARIAN ---
st.subheader("2. Kegiatan Harian")
if 'kegiatan_count' not in st.session_state: st.session_state.kegiatan_count = 1

kegiatan_list = []
for i in range(st.session_state.kegiatan_count):
    col_w, col_u, col_k, col_s = st.columns([1.5, 3, 2, 1])
    with col_w: wkt = st.text_input(f"Waktu", "07.30-14.30", key=f"w{i}", label_visibility="collapsed" if i>0 else "visible")
    with col_u: uraian = st.text_input(f"Uraian", "Stand By On Site dan Monitoring IP UNEJ", key=f"u{i}", label_visibility="collapsed" if i>0 else "visible")
    with col_k: ket = st.text_input(f"Keterangan", "-", key=f"k{i}", label_visibility="collapsed" if i>0 else "visible")
    with col_s: stat = st.selectbox(f"Status", ["Selesai", "Proses", "Pending"], key=f"s{i}", label_visibility="collapsed" if i>0 else "visible")
    kegiatan_list.append({'waktu': wkt, 'uraian': uraian, 'ket': ket, 'status': stat})

if st.button("➕ Tambah Kegiatan"):
    st.session_state.kegiatan_count += 1
    st.rerun()

# --- FORM KENDALA & FOLLOW UP ---
st.divider()
col_kendala, col_followup = st.columns(2)

with col_kendala:
    with st.expander("⚠️ 3. Kendala / Insiden"):
        if 'kendala_count' not in st.session_state: st.session_state.kendala_count = 1
        kendala_list = []
        for i in range(st.session_state.kendala_count):
            st.markdown(f"**Kendala {i+1}**")
            m = st.text_input("Masalah", "-", key=f"km{i}")
            t = st.text_input("Tindakan", "-", key=f"kt{i}")
            h = st.text_input("Hasil", "-", key=f"kh{i}")
            kendala_list.append({'masalah': m, 'tindakan': t, 'hasil': h})
        if st.button("➕ Tambah Kendala"):
            st.session_state.kendala_count += 1
            st.rerun()

with col_followup:
    with st.expander("📌 4. Follow Up"):
        if 'fu_count' not in st.session_state: st.session_state.fu_count = 1
        followup_list = []
        for i in range(st.session_state.fu_count):
            st.markdown(f"**Follow Up {i+1}**")
            d = st.text_input("Deskripsi", "-", key=f"fd{i}")
            a = st.text_input("Alasan/Kendala", "-", key=f"fa{i}")
            t = st.text_input("Target", "-", key=f"ft{i}")
            followup_list.append({'deskripsi': d, 'alasan': a, 'target': t})
        if st.button("➕ Tambah Follow Up"):
            st.session_state.fu_count += 1
            st.rerun()

# --- FORM FOTO ---
st.divider()
st.subheader("📸 Lampiran Status IP")
locations = ["SPTN", "Pasuruan", "Lumajang", "Jubung", "FKIP2", "Bondowoso", "Bhayangkara", "IDREN UB"]
ip_photos_data = []

cols = st.columns(4) 
for idx, loc in enumerate(locations):
    with cols[idx % 4]:
        with st.container(border=True):
            st.markdown(f"**IP {loc}**")
            up_file = st.file_uploader("Upload", type=['png', 'jpg', 'jpeg'], key=f"file_{loc}", label_visibility="collapsed")
            status_toggle = st.toggle("UP", value=True, key=f"tog_{loc}")
            status_text = "Up" if status_toggle else "Down"
            ip_photos_data.append({'name': loc, 'file': up_file, 'status': status_text})

st.divider()
if st.button("🚀 Generate Report", type="primary", use_container_width=True):
    # Pass teks_tgl_indo ke fungsi generate
    final_docx = generate_report(file_template, teks_tgl_indo, lok, kegiatan_list, kendala_list, followup_list, ip_photos_data)
    
    if final_docx:
        st.success(f"Laporan berhasil dibuat!")
        
        # Format nama file otomatis sesuai dengan yang kamu inginkan
        nama_file_download = f"Daily Report_Managed Service_Universitas Jember_{teks_tgl_indo}.docx"
        
        st.download_button(
            label=f"⬇️ Download {nama_file_download}",
            data=final_docx.getvalue(),
            file_name=nama_file_download,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
